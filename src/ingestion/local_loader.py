"""
本地文件夹加载模块

功能：
- 扫描本地目录，加载所有代码和文档文件
- 支持 PDF、Word、Markdown、纯文本
- 与 GitHubDocument 格式统一，方便后续统一处理
"""

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from src.config import config

logger = logging.getLogger(__name__)


@dataclass
class LocalDocument:
    """本地加载的文档"""
    source: str          # 来源标识 e.g. "local:/path/to/file"
    file_path: str       # 绝对路径
    content: str
    file_type: str       # "code" | "doc"
    language: str
    size: int
    metadata: dict = field(default_factory=dict)


class LocalLoader:
    """加载本地文件夹中的文档"""

    def load_directory(
        self,
        directory: str,
        recursive: bool = True,
        file_filter: Optional[list[str]] = None,
    ) -> list[LocalDocument]:
        """
        Args:
            directory:    目录路径
            recursive:    是否递归子目录
            file_filter:  文件扩展名白名单，None 表示全部支持的类型
        """
        dir_path = Path(directory).resolve()
        if not dir_path.is_dir():
            raise NotADirectoryError(f"路径不存在或不是目录: {directory}")

        pattern = "**/*" if recursive else "*"
        docs = []
        all_known = config.CODE_EXTENSIONS | config.DOC_EXTENSIONS

        for fp in dir_path.glob(pattern):
            if not fp.is_file():
                continue
            ext = fp.suffix.lower()

            # 过滤
            if file_filter and ext not in file_filter:
                continue
            if ext not in all_known:
                continue
            # 跳过隐藏文件
            if any(part.startswith(".") for part in fp.parts):
                continue

            content = self._read_file(fp)
            if content is None:
                continue

            file_type = "code" if ext in config.CODE_EXTENSIONS else "doc"
            from src.ingestion.github_loader import EXT_TO_LANGUAGE
            language = EXT_TO_LANGUAGE.get(ext, "unknown")

            docs.append(LocalDocument(
                source=f"local:{fp}",
                file_path=str(fp),
                content=content,
                file_type=file_type,
                language=language,
                size=fp.stat().st_size,
                metadata={"directory": str(dir_path)},
            ))

        logger.info(f"本地加载完成: {len(docs)} 个文件 from {directory}")
        return docs

    def load_file(self, file_path: str) -> Optional[LocalDocument]:
        """加载单个文件"""
        fp = Path(file_path).resolve()
        if not fp.is_file():
            return None
        content = self._read_file(fp)
        if content is None:
            return None
        ext = fp.suffix.lower()
        from src.ingestion.github_loader import EXT_TO_LANGUAGE
        return LocalDocument(
            source=f"local:{fp}",
            file_path=str(fp),
            content=content,
            file_type="code" if ext in config.CODE_EXTENSIONS else "doc",
            language=EXT_TO_LANGUAGE.get(ext, "unknown"),
            size=fp.stat().st_size,
            metadata={},
        )

    @staticmethod
    def _read_file(fp: Path) -> Optional[str]:
        ext = fp.suffix.lower()
        try:
            if ext == ".pdf":
                return LocalLoader._read_pdf(fp)
            elif ext == ".docx":
                return LocalLoader._read_docx(fp)
            else:
                return fp.read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            logger.warning(f"读取文件失败 {fp}: {e}")
            return None

    @staticmethod
    def _read_pdf(fp: Path) -> str:
        from pypdf import PdfReader
        reader = PdfReader(str(fp))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    @staticmethod
    def _read_docx(fp: Path) -> str:
        from docx import Document
        doc = Document(str(fp))
        return "\n".join(para.text for para in doc.paragraphs)
